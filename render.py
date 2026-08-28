"""content.json → md / html / pdf. 양식은 박제. 모델 호출 없음."""
from __future__ import annotations

import html
import json
import sys
from pathlib import Path

CSS = """
@page { size: A4; margin: 14mm 15mm 14mm 15mm; }
* { box-sizing: border-box; }
body { font-family:'NanumGothic','Apple SD Gothic Neo','AppleSDGothicNeo',sans-serif;
       word-break: keep-all; line-height: 1.5; color:#111; margin:0; }
.title { text-align:center; font-weight:800; font-size:17pt; letter-spacing:1px;
         margin:0 0 9px; }
.sheet { border:1px solid #000; }
table.hdr { width:100%; border-collapse:collapse; table-layout:fixed; }
table.hdr td { border:1px solid #000; padding:5px 8px; font-size:10pt; vertical-align:middle; }
table.hdr tr:first-child td { border-top:0; }
table.hdr td:first-child { border-left:0; }
table.hdr td:last-child { border-right:0; }
td.lab { background:#EAEDF2; text-align:center; font-weight:700; white-space:nowrap; }
.band { text-align:center; font-weight:800; font-size:10.5pt; padding:5px 0;
        border-top:1px solid #000; border-bottom:1px solid #000; }
.body { padding:9px 11px 12px; }
.sec { font-weight:800; font-size:10.5pt; margin:9px 0 3px; }
.sec:first-child { margin-top:2px; }
.lead { margin:2px 0 4px; }
.li { padding-left:12px; text-indent:-12px; margin:2px 0; font-size:10pt; }
.li.red { color:#C0392B; }
.sub { padding-left:24px; text-indent:-12px; margin:2px 0; font-size:10pt; }
ol.plan { margin:2px 0; padding-left:20px; }
ol.plan li { margin:2px 0; font-size:10pt; }
.para { margin:3px 0; font-size:10pt; }
table.ct { width:100%; border-collapse:collapse; table-layout:fixed; margin:5px 0 4px; }
table.ct th, table.ct td { border:1px solid #000; padding:4px 6px; font-size:9.5pt;
                           vertical-align:top; word-break:keep-all; }
table.ct th { background:#EAEDF2; font-weight:700; text-align:center; }
table.ct td.k { font-weight:700; text-align:center; }
.note { font-size:8.5pt; color:#555; margin:4px 2px 0; line-height:1.45; }
.closing { text-align:right; font-size:10pt; margin:8px 6px 2px; }
"""

FOOTER = '<div style="width:100%;text-align:center;font-size:8pt;color:#333;">- <span class="pageNumber"></span> -</div>'

FONT_CANDIDATES = [
    Path(__file__).resolve().parent / "fonts" / "NanumGothic-Regular.ttf",
    Path.home() / "Library" / "Fonts" / "NanumGothic.ttf",
    Path("C:/Windows/Fonts/malgun.ttf"),
]


def esc(s) -> str:
    return html.escape(str(s), quote=False)


def render_md(d: dict) -> str:
    lines = [f"# {d['title']}", ""]
    hdr = d.get("header") or []
    if hdr:
        lines.append("| 항목 | 내용 |")
        lines.append("|---|---|")
        for row in hdr:
            if len(row) == 4:
                lines.append(f"| {row[0]} | {row[1]} |")
                lines.append(f"| {row[2]} | {row[3]} |")
            elif len(row) >= 2:
                lines.append(f"| {row[0]} | {row[1]} |")
        lines.append("")
    lines.append("## < 회 의 내 용 >")
    lines.append("")
    for sec in d.get("sections") or []:
        lines.append(f"### {sec.get('heading','')}")
        lines.extend(_md_blocks(sec.get("blocks") or []))
        lines.append("")
    lines.append(d.get("closing", "끝."))
    ap = d.get("appendix")
    if ap:
        lines.append("")
        lines.append(f"## {ap.get('band','부록')}")
        if ap.get("note"):
            lines.append(ap["note"])
        lines.extend(_md_blocks(ap.get("blocks") or []))
        if ap.get("table"):
            lines.extend(_md_table(ap["table"]))
    return "\n".join(lines).rstrip() + "\n"


def _md_blocks(blocks: list) -> list[str]:
    out: list[str] = []
    for b in blocks:
        if b.get("lead"):
            out.append(f"- {b['lead']}")
        for it in b.get("bullets") or []:
            if isinstance(it, dict):
                out.append(f"- {it.get('li','')}")
                for s in it.get("sub") or []:
                    out.append(f"  - {s}")
            else:
                out.append(f"- {it}")
        for i, x in enumerate(b.get("ordered") or [], 1):
            out.append(f"{i}. {x}")
        if b.get("para"):
            out.append(b["para"])
        if b.get("table"):
            out.extend(_md_table(b["table"]))
    return out


def _md_table(t: dict) -> list[str]:
    headers = t.get("headers") or []
    rows = t.get("rows") or []
    if not headers:
        return []
    out = ["", "| " + " | ".join(headers) + " |", "|" + "|".join("---" for _ in headers) + "|"]
    for r in rows:
        cells = [str(c) for c in r] + [""] * max(0, len(headers) - len(r))
        out.append("| " + " | ".join(cells[: len(headers)]) + " |")
    out.append("")
    return out


def build_html(d: dict) -> str:
    secs = []
    for s in d.get("sections", []):
        secs.append(f'<div class="sec">{esc(s["heading"])}</div>')
        secs.append(_html_blocks(s.get("blocks", [])))
    note = f'<div class="note">{esc(d["note"])}</div>' if d.get("note") else ""
    closing = f'<div class="closing">{esc(d.get("closing", "끝."))}</div>'
    appendix = ""
    ap = d.get("appendix")
    if ap:
        apnote = f'<div class="note">{esc(ap["note"])}</div>' if ap.get("note") else ""
        aptab = _html_table(ap["table"]) if ap.get("table") else ""
        apblocks = _html_blocks(ap.get("blocks", [])) if ap.get("blocks") else ""
        appendix = (
            f'<div class="band">{esc(ap.get("band", "< 부록 >"))}</div>'
            f'<div class="body">{apnote}{aptab}{apblocks}</div>'
        )
    face = _font_face()
    return f"""<!doctype html><html><head><meta charset="utf-8"><style>{face}{CSS}</style></head>
<body>
<div class="title">{esc(d["title"])}</div>
<div class="sheet">
  <table class="hdr">{_html_header(d.get("header") or [])}</table>
  {note}
  <div class="band">&lt; 회 의 내 용 &gt;</div>
  <div class="body">{''.join(secs)}{closing}</div>
  {appendix}
</div>
</body></html>"""


def html_to_pdf(html_str: str, pdf_path: Path) -> None:
    from playwright.sync_api import sync_playwright

    pdf_path = Path(pdf_path)
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_str, wait_until="load")
        page.pdf(
            path=str(pdf_path),
            format="A4",
            print_background=True,
            display_header_footer=True,
            header_template="<span></span>",
            footer_template=FOOTER,
            margin={"top": "14mm", "bottom": "16mm", "left": "15mm", "right": "15mm"},
        )
        browser.close()


def write_outputs(
    d: dict,
    out_dir: Path,
    stem: str,
    *,
    pdf: bool = False,
    extras: bool = False,
) -> dict[str, Path]:
    """최종본은 docx. json은 재생성용. pdf/html/md는 extras일 때만."""
    out_dir.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    docx_path = out_dir / f"{stem}.docx"
    build_docx(d, docx_path)
    paths["docx"] = docx_path
    js = out_dir / f"{stem}.json"
    js.write_text(json.dumps(d, ensure_ascii=False, indent=2), encoding="utf-8")
    paths["json"] = js
    if extras:
        md = out_dir / f"{stem}.md"
        ht = out_dir / f"{stem}.html"
        md.write_text(render_md(d), encoding="utf-8")
        html_str = build_html(d)
        ht.write_text(html_str, encoding="utf-8")
        paths["md"] = md
        paths["html"] = ht
        if pdf:
            pdf_path = out_dir / f"{stem}.pdf"
            html_to_pdf(html_str, pdf_path)
            paths["pdf"] = pdf_path
    elif pdf:
        html_str = build_html(d)
        pdf_path = out_dir / f"{stem}.pdf"
        html_to_pdf(html_str, pdf_path)
        paths["pdf"] = pdf_path
    return paths


def build_docx(d: dict, path: Path) -> Path:
    """개조식 워드. 테두리 시트 없이 제목·메타·불릿·표만."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    font = "Malgun Gothic" if sys.platform.startswith("win") else "NanumGothic"
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font)
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.35

    def run_font(run, size=11, bold=False, color=None):
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            rf.set(qn(a), font)

    def p(text="", *, size=11, bold=False, align=None, space_before=0, space_after=4, color=None):
        para = doc.add_paragraph()
        if align:
            para.alignment = align
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.line_spacing = 1.35
        if text:
            r = para.add_run(text)
            run_font(r, size=size, bold=bold, color=color)
        return para

    def bullet(text, *, indent=False, color=None):
        para = p("", size=11, space_after=3)
        para.paragraph_format.left_indent = Cm(0.9 if indent else 0.45)
        para.paragraph_format.first_line_indent = Cm(-0.4)
        r = para.add_run(("  ·  " if indent else "·  ") + text)
        run_font(r, size=11, color=color)

    p(d.get("title") or "보고서", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
      space_before=0, space_after=12)

    for label, value in _header_pairs(d.get("header") or []):
        if _blank(value):
            continue
        para = p("", space_after=2)
        r1 = para.add_run(f"{label}  ")
        run_font(r1, size=10.5, bold=True, color="555555")
        r2 = para.add_run(value)
        run_font(r2, size=11)

    if d.get("note") and not _blank(d["note"]):
        p(d["note"], size=9, color="666666", space_before=6, space_after=4)

    p("회의 내용", size=12, bold=True, space_before=14, space_after=8)

    for sec in d.get("sections") or []:
        heading = str(sec.get("heading") or "").strip()
        if heading:
            p(heading, size=12, bold=True, space_before=12, space_after=6)
        _docx_blocks(sec.get("blocks") or [], p, bullet, run_font, doc, font)

    p(d.get("closing", "끝."), size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=14, space_after=0)

    ap = d.get("appendix")
    if ap:
        p(str(ap.get("band") or "부록").strip("<> "), size=12, bold=True, space_before=16, space_after=6)
        if ap.get("note"):
            p(ap["note"], size=9, color="666666")
        _docx_blocks(ap.get("blocks") or [], p, bullet, run_font, doc, font)
        if ap.get("table"):
            _docx_table(doc, ap["table"], run_font)

    path = Path(path)
    doc.save(str(path))
    return path


def _blank(s: str) -> bool:
    t = (s or "").strip()
    return (not t) or t in {"(원문에 없음)", "-", "—"}


def _header_pairs(rows) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    for row in rows:
        if len(row) == 4:
            pairs.append((str(row[0]).strip(), str(row[1]).strip()))
            pairs.append((str(row[2]).strip(), str(row[3]).strip()))
        elif len(row) >= 2:
            pairs.append((str(row[0]).strip(), str(row[1]).strip()))
    return pairs


def _docx_blocks(blocks, p, bullet, run_font, doc, font):
    from docx.shared import Cm

    for b in blocks:
        if b.get("lead"):
            bullet(str(b["lead"]))
        for it in b.get("bullets") or []:
            if isinstance(it, dict):
                color = "C0392B" if it.get("color") == "red" else None
                bullet(str(it.get("li") or ""), color=color)
                for s in it.get("sub") or []:
                    bullet(str(s), indent=True)
            else:
                bullet(str(it))
        for i, x in enumerate(b.get("ordered") or [], 1):
            para = p("", space_after=3)
            para.paragraph_format.left_indent = Cm(0.5)
            r = para.add_run(f"{i}.  {x}")
            run_font(r, size=11)
        if b.get("para"):
            p(str(b["para"]), space_after=6)
        if b.get("table"):
            _docx_table(doc, b["table"], run_font)


def _docx_table(doc, t: dict, run_font) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    headers = t.get("headers") or []
    rows = t.get("rows") or []
    if not headers:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def cell_run(cell, text, *, bold=False, center=False):
        para = cell.paragraphs[0]
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(str(text))
        run_font(r, size=10, bold=bold)

    for c, h in enumerate(headers):
        shade(tbl.rows[0].cells[c], "F2F2F2")
        cell_run(tbl.rows[0].cells[c], h, bold=True, center=True)
    fcb = t.get("firstcol_bold")
    for ri, rv in enumerate(rows, 1):
        for c, h in enumerate(headers):
            val = rv[c] if c < len(rv) else ""
            cell_run(tbl.rows[ri].cells[c], val, bold=bool(fcb and c == 0), center=bool(fcb and c == 0))

    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "BBBBBB")
        borders.append(e)
    tblPr.append(borders)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def _font_face() -> str:
    for p in FONT_CANDIDATES:
        if p.exists():
            uri = p.resolve().as_uri()
            return f"@font-face {{ font-family:'NanumGothic'; src:url('{uri}'); }}\n"
    return ""


def _html_header(rows) -> str:
    out = []
    for i, r in enumerate(rows):
        if i == 0 and len(r) == 4:
            out.append(
                f'<tr><td class="lab" style="width:15%">{esc(r[0])}</td>'
                f'<td style="width:37%">{esc(r[1])}</td>'
                f'<td class="lab" style="width:13%">{esc(r[2])}</td>'
                f'<td style="width:35%">{esc(r[3])}</td></tr>'
            )
        elif len(r) >= 2:
            out.append(
                f'<tr><td class="lab">{esc(r[0])}</td>'
                f'<td colspan="3">{esc(r[1])}</td></tr>'
            )
    return "".join(out)


def _html_blocks(blocks) -> str:
    out = []
    for b in blocks:
        if "lead" in b:
            out.append(f'<div class="lead li">- {esc(b["lead"])}</div>')
        if "bullets" in b:
            out.append(_html_bullets(b["bullets"]))
        if "ordered" in b:
            lis = "".join(f"<li>{esc(x)}</li>" for x in b["ordered"])
            out.append(f'<ol class="plan">{lis}</ol>')
        if "para" in b:
            out.append(f'<div class="para">{esc(b["para"])}</div>')
        if "table" in b:
            out.append(_html_table(b["table"]))
    return "".join(out)


def _html_bullets(items) -> str:
    out = []
    for it in items:
        if isinstance(it, dict):
            cls = "li red" if it.get("color") == "red" else "li"
            out.append(f'<div class="{cls}">- {esc(it.get("li",""))}</div>')
            for s in it.get("sub", []):
                out.append(f'<div class="sub">{esc(s)}</div>')
        else:
            out.append(f'<div class="li">- {esc(it)}</div>')
    return "".join(out)


def _html_table(t) -> str:
    widths = t.get("widths")
    colg = "".join(f'<col style="width:{w}%">' for w in widths) if widths else ""
    heads = "".join(f"<th>{esc(h)}</th>" for h in t["headers"])
    fcb = t.get("firstcol_bold")
    rows = []
    for rv in t["rows"]:
        cells = []
        for c, v in enumerate(rv):
            klass = ' class="k"' if (fcb and c == 0) else ""
            cells.append(f"<td{klass}>{esc(v)}</td>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return (
        f'<table class="ct"><colgroup>{colg}</colgroup>'
        f"<thead><tr>{heads}</tr></thead><tbody>{''.join(rows)}</tbody></table>"
    )
